# -*- coding: utf-8 -*-
"""Compile MoonEvidence项目报告.md -> .docx via pandoc + the shared
Chinese-academic reference doc, then post-process:
  1) tables   -> centered, three-line (booktabs) style
  2) formulas -> display equations centered with right-aligned (N) numbers
  3) code     -> boxed (paragraph borders merge into one frame)
  4) body     -> first-line indent of 2 chars
"""
import os
import subprocess
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _set_first_line_indent(p, chars=200):
    """首行缩进 2 个字符（firstLineChars=200，与字号无关）。"""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars))
    ind.set(qn("w:firstLine"), "480")


def _is_body(p, eq_paras):
    """正文段判定；展示公式段（eq_paras）不缩进，行内公式段照常缩进。"""
    if p._p in eq_paras:
        return False
    name = (p.style.name or "").lower()
    if name.startswith("heading") or name == "title":
        return False
    if any(k in name for k in ("caption", "toc", "source", "verbatim", "code", "list", "quote")):
        return False
    if p._p.findall(".//" + qn("w:drawing")):
        return False
    return bool(p.text.strip())


def _border(sz, space="0"):
    el_attrs = {"w:val": "single", "w:sz": str(sz), "w:space": space, "w:color": "000000"}
    return el_attrs


def _mk(tag, attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    return el


def _three_line_table(t):
    """三线表：仅保留顶线、表头下线、底线，其余边框清空。"""
    tbl_pr = t._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    borders.append(_mk("w:top", _border(12)))
    borders.append(_mk("w:bottom", _border(12)))
    for side in ("w:left", "w:right", "w:insideH", "w:insideV"):
        borders.append(_mk(side, {"w:val": "none", "w:sz": "0", "w:space": "0", "w:color": "auto"}))
    tbl_pr.append(borders)
    # 清掉单元格残留边框，再给表头行加下细线
    for ri, row in enumerate(t.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            for old in tc_pr.findall(qn("w:tcBorders")):
                tc_pr.remove(old)
            if ri == 0:
                tcb = OxmlElement("w:tcBorders")
                tcb.append(_mk("w:bottom", _border(6)))
                tc_pr.append(tcb)


def _page_content_width_tw(doc):
    """正文可用宽度（twips）；sectPr 缺省时按 A4 + 1 英寸页边距计。"""
    sect = doc.sections[0]
    if sect.page_width and sect.left_margin is not None and sect.right_margin is not None:
        return int(sect.page_width - sect.left_margin - sect.right_margin) // 635
    return 11906 - 1440 * 2


def _number_equations(doc):
    """公式段居中排版并在行末右对齐追加 (N) 编号；返回 (数量, 公式段集合)。"""
    content_tw = _page_content_width_tw(doc)
    center_pos = content_tw // 2
    right_pos = content_tw
    n = 0
    eq_paras = set()
    for p in doc.paragraphs:
        omath_paras = p._p.findall(".//{%s}oMathPara" % M_NS)
        if not omath_paras:
            continue
        n += 1
        eq_paras.add(p._p)
        omp = omath_paras[0]
        omath = omp.find("{%s}oMath" % M_NS)
        parent = omp.getparent()
        # 块级公式转为行内元素，用制表位实现「居中公式 + 右侧编号」
        parent.replace(omp, omath)
        pPr = p._p.get_or_add_pPr()
        for old in pPr.findall(qn("w:tabs")):
            pPr.remove(old)
        tabs = OxmlElement("w:tabs")
        tabs.append(_mk("w:tab", {"w:val": "center", "w:pos": str(center_pos)}))
        tabs.append(_mk("w:tab", {"w:val": "right", "w:pos": str(right_pos)}))
        pPr.append(tabs)
        for old in pPr.findall(qn("w:jc")):
            pPr.remove(old)
        # 公式前后各一个制表符，编号推到右缘
        tab_before = OxmlElement("w:r")
        tab_before.append(OxmlElement("w:tab"))
        omath.addprevious(tab_before)
        tab_after = OxmlElement("w:r")
        tab_after.append(OxmlElement("w:tab"))
        num_run = OxmlElement("w:r")
        num_text = OxmlElement("w:t")
        num_text.text = f"({n})"
        num_run.append(num_text)
        omath.addnext(num_run)
        omath.addnext(tab_after)
    return n, eq_paras


def _force_mono_font(pel, size_half_pt="18"):
    """代码段内所有 run 强制 Consolas（中文回落宋体）与统一字号。"""
    for r in pel.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r.insert(0, rPr)
        for old in rPr.findall(qn("w:rFonts")):
            rPr.remove(old)
        rPr.insert(0, _mk("w:rFonts", {
            "w:ascii": "Consolas", "w:hAnsi": "Consolas",
            "w:cs": "Consolas", "w:eastAsia": "宋体",
        }))
        for tag in ("w:sz", "w:szCs"):
            for old in rPr.findall(qn(tag)):
                rPr.remove(old)
            rPr.append(_mk(tag, {"w:val": size_half_pt}))


def _box_code_blocks(doc, content_tw):
    """代码块包进单行单列表格：跨页时每页边框闭合，比段落边框可靠。"""
    code_paras = []
    for p in doc.paragraphs:
        name = (p.style.name or "").lower()
        if any(k in name for k in ("source", "verbatim")):
            code_paras.append(p._p)
    for pel in code_paras:
        _force_mono_font(pel)
        parent = pel.getparent()

        tbl = OxmlElement("w:tbl")
        tblPr = OxmlElement("w:tblPr")
        tblPr.append(_mk("w:tblW", {"w:w": str(content_tw), "w:type": "dxa"}))
        borders = OxmlElement("w:tblBorders")
        for side in ("w:top", "w:left", "w:bottom", "w:right"):
            borders.append(_mk(side, {"w:val": "single", "w:sz": "6", "w:space": "0", "w:color": "7F7F7F"}))
        tblPr.append(borders)
        margins = OxmlElement("w:tblCellMar")
        margins.append(_mk("w:top", {"w:w": "80", "w:type": "dxa"}))
        margins.append(_mk("w:bottom", {"w:w": "80", "w:type": "dxa"}))
        margins.append(_mk("w:left", {"w:w": "140", "w:type": "dxa"}))
        margins.append(_mk("w:right", {"w:w": "140", "w:type": "dxa"}))
        tblPr.append(margins)
        tbl.append(tblPr)
        grid = OxmlElement("w:tblGrid")
        grid.append(_mk("w:gridCol", {"w:w": str(content_tw)}))
        tbl.append(grid)
        tr = OxmlElement("w:tr")
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        tcPr.append(_mk("w:tcW", {"w:w": str(content_tw), "w:type": "dxa"}))
        tcPr.append(_mk("w:shd", {"w:val": "clear", "w:color": "auto", "w:fill": "F7F7F7"}))
        tc.append(tcPr)
        tr.append(tc)
        tbl.append(tr)

        parent.replace(pel, tbl)
        tc.append(pel)
    return len(code_paras)


sys.stdout.reconfigure(encoding="utf-8")

HERE = os.path.dirname(os.path.abspath(__file__))
REF = r"E:\大学\交作业\区块链技术与应用\labs\_ref.docx"
MD = "MoonEvidence项目报告.md"
DOCX = "MoonEvidence项目报告.docx"

subprocess.run(
    ["pandoc", MD, "-o", DOCX, "--reference-doc=" + REF],
    cwd=HERE, check=True,
)

path = os.path.join(HERE, DOCX)
d = Document(path)

for t in d.tables:
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _three_line_table(t)
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _keep_table_together(t):
    """整表不跨页：除末行外全部段落 keepNext，行内禁止拆分。"""
    rows = t.rows
    for ri, row in enumerate(rows):
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))
        if ri == len(rows) - 1:
            continue
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._p.get_or_add_pPr()
                if pPr.find(qn("w:keepNext")) is None:
                    pPr.append(OxmlElement("w:keepNext"))


for t in d.tables:
    if any("陈俊文" in c.text for row in t.rows for c in row.cells):
        _keep_table_together(t)

eq, eq_paras = _number_equations(d)
code = _box_code_blocks(d, _page_content_width_tw(d))

indented = 0
for p in d.paragraphs:
    if _is_body(p, eq_paras):
        _set_first_line_indent(p)
        indented += 1

d.save(path)

h2 = sum(1 for p in d.paragraphs if p.style.name == "Heading 2")
print(f"{DOCX} | tables={len(d.tables)}(三线) | 公式编号={eq} | 代码框段={code} | H2={h2} | 缩进段={indented}")
print("done")
